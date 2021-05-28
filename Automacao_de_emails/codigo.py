import pandas as pd
import docx

from docx import Document
from docx.shared import Inches, Cm, Pt, Length
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date


df = pd.read_excel('random_data.xlsx')  # ler planilha de dados

i = df[df.columns[0]].count()   - 1

while i >= 0:

    # a parte de baixo separa as informações em cada variável que estão relacionadas ao cod. do produto
    cliente_nome = df.iat[i,0]
    cliente_end = df.iat[i,1]
    cliente_cep = df.iat[i,2]
    data_entrega = df.iat[i,3].strftime('%d/%m/%Y')
    num_OC = df.iat[i,4]
    produto_nome = df.iat[i,5]



    arquivo =open('docref.docx','rb') # Abrir arquivo de referÊncia para escrever
    document =Document(arquivo)
    paragraph = document.add_paragraph()   

    # Iniciar escrita
    paragraph = document.add_paragraph('Caro(a) ' + cliente_nome + ':')

    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing
    paragraph_format.line_spacing_rule
    paragraph_format.line_spacing = Pt(12)    #espaçamento de 12 pt

    paragraph = document.add_paragraph('Notificamos que na data ' + str(data_entrega) + ', tentamos entregar seu produto ' 
                       + produto_nome + ' com o número de referência ' + str(num_OC) + ' no endereço ' 
                       + cliente_end + ' de CEP ' + str(cliente_cep)+ '. Deve ter havido um mal-entendido em relação' 
                       ' a esta entrega porque não havia ninguém disponível para assinar o pedido na data de entrega'
                       ' agendada, portanto não entregamos o pedido. Pedimos que entre em contato assim que for possível'
                       'para que possamos agendar outra hora de entrega.')

    paragraph.alignment= WD_ALIGN_PARAGRAPH.JUSTIFY    #Justificar

    paragraph_format.line_spacing = Pt(12)   #espaçamento de 12 pt
    paragraph = document.add_paragraph('Mais uma vez, obrigado pelo pedido. Esperamos fazer sua entrega em breve' 
                       'e atender às suas necessidades no futuro.')

    paragraph.alignment= WD_ALIGN_PARAGRAPH.JUSTIFY    #Justificar

    paragraph_format.line_spacing = Pt(12)     #espaçamento de 12 pt
    paragraph = document.add_paragraph('Atenciosamente,')
    paragraph_format.line_spacing = Pt(12)
    paragraph = document.add_paragraph('Vert TECH')   #espaçamento de 12 pt
    
    document.save(cliente_nome + '_email.docx') #salvar documento
    i-=1
