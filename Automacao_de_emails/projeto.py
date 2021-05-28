import pandas as pd
import docx
from docx import Document
from docx.shared import Inches, Cm, Pt, Length
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date



df = pd.read_excel('random_data.xlsx')  # ler planilha de dados
merc_ref = [input()] # recebe o numero de referência da mercadoria

# a parte de baixo separa as informações em cada variável que estão relacionadas ao cod. do produto
cliente_nome = df[df.NUMERO_OC.isin(merc_ref)].iat[0,0]
cliente_end = df[df.NUMERO_OC.isin(merc_ref)].iat[0,1]
cliente_cep = df[df.NUMERO_OC.isin(merc_ref)].iat[0,2]
data_entrega = df[df.NUMERO_OC.isin(merc_ref)].iat[0,3].strftime('%d/%m/%Y')
num_OC = df[df.NUMERO_OC.isin(merc_ref)].iat[0,4]
produto_nome = df[df.NUMERO_OC.isin(merc_ref)].iat[0,5]

arquivo =open('docref.docx','rb')
document =Document(arquivo)
paragraph = document.add_paragraph()
paragraph = document.add_paragraph('Caro(a) ' + cliente_nome + ':')

paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing
paragraph_format.line_spacing_rule
paragraph_format.line_spacing = Pt(12)

paragraph = document.add_paragraph('Notificamos que na data ' + str(data_entrega) + ', tentamos entregar seu produto ' 
                       + produto_nome + ' com o número de referência ' + str(num_OC) + ' no endereço ' 
                       + cliente_end + ' de CEP ' + str(cliente_cep)+ '. Deve ter havido um mal-entendido em relação' 
                       ' a esta entrega porque não havia ninguém disponível para assinar o pedido na data de entrega'
                       ' agendada, portanto não entregamos o pedido. Pedimos que entre em contato assim que for possível'
                       'para que possamos agendar outra hora de entrega.')

paragraph.alignment= WD_ALIGN_PARAGRAPH.JUSTIFY

paragraph_format.line_spacing = Pt(12)
paragraph = document.add_paragraph('Mais uma vez, obrigado pelo pedido. Esperamos fazer sua entrega em breve' 
                       'e atender às suas necessidades no futuro.')

paragraph.alignment= WD_ALIGN_PARAGRAPH.JUSTIFY

paragraph_format.line_spacing = Pt(12)
paragraph = document.add_paragraph('Atenciosamente,')
paragraph_format.line_spacing = Pt(12)
paragraph = document.add_paragraph('Vert TECH')


document.save('email.docx') #salvar documento
